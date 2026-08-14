import os
import re
from io import BytesIO
from urllib.parse import urlparse

import requests
import PyPDF2
from docx import Document


COMMON_SKILLS = [
    "python",
    "django",
    "react",
    "javascript",
    "html",
    "css",
    "sql",
    "mysql",
    "postgresql",
    "git",
    "github",
    "docker",
    "aws",
    "rest api",
]


def extract_resume_text(file_source):
    text = ""

    # Cloudinary URL
    if isinstance(file_source, str) and file_source.startswith(
        ("http://", "https://")
    ):

        response = requests.get(
            file_source,
            timeout=30,
        )

        response.raise_for_status()

        file_data = BytesIO(response.content)

        path = urlparse(file_source).path
        extension = os.path.splitext(path)[1].lower()

    # Local file path
    else:

        extension = os.path.splitext(file_source)[1].lower()
        file_data = file_source

    if extension == ".pdf":

        reader = PyPDF2.PdfReader(file_data)

        for page in reader.pages:

            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    elif extension == ".docx":

        document = Document(file_data)

        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"

    return text.strip()


def analyze_resume(text):

    text_lower = text.lower()

    found_skills = []

    for skill in COMMON_SKILLS:

        if skill in text_lower:
            found_skills.append(skill)

    missing_skills = [
        skill
        for skill in COMMON_SKILLS
        if skill not in found_skills
    ]

    score = 50

    score += min(len(found_skills) * 3, 30)

    if len(text.split()) > 300:
        score += 10

    email_found = bool(
        re.search(
            r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
            text,
        )
    )

    phone_found = bool(
        re.search(
            r"\+?\d[\d\s-]{8,}",
            text,
        )
    )

    if email_found:
        score += 5

    if phone_found:
        score += 5

    score = min(score, 100)

    suggestions = []

    if len(found_skills) < 6:
        suggestions.append(
            "Add more relevant technical skills."
        )

    if not email_found:
        suggestions.append(
            "Add a professional email address."
        )

    if not phone_found:
        suggestions.append(
            "Add a phone number."
        )

    if len(text.split()) < 300:
        suggestions.append(
            "Add more projects, internships, and measurable achievements."
        )

    return {
        "ats_score": score,
        "skills_found": found_skills,
        "missing_skills": missing_skills,
        "suggestions": suggestions,
    }
import os
import re
from io import BytesIO
from urllib.parse import urlparse

import requests
import PyPDF2
from docx import Document


COMMON_SKILLS = [
    "python",
    "django",
    "react",
    "javascript",
    "html",
    "css",
    "sql",
    "mysql",
    "postgresql",
    "git",
    "github",
    "docker",
    "aws",
    "rest api",
    "flask",
    "bootstrap",
    "tailwind",
    "mongodb",
]


def extract_resume_text(file_source):
    """
    Extract text from PDF or DOCX resume.

    Supports:
    - Cloudinary URLs
    - Local file paths
    - Cloudinary URLs where file extension is not obvious
    """

    text = ""
    file_data = None
    extension = ""

    # ==========================================
    # CLOUDINARY / REMOTE FILE
    # ==========================================

    if isinstance(file_source, str) and file_source.startswith(
        ("http://", "https://")
    ):

        response = requests.get(
            file_source,
            timeout=30,
        )

        response.raise_for_status()

        file_data = BytesIO(response.content)

        # First try URL extension
        path = urlparse(file_source).path
        extension = os.path.splitext(path)[1].lower()

        # If URL has no extension, use Content-Type
        if not extension:

            content_type = response.headers.get(
                "Content-Type",
                ""
            ).lower()

            if "pdf" in content_type:
                extension = ".pdf"

            elif (
                "word" in content_type
                or "document" in content_type
                or "docx" in content_type
            ):
                extension = ".docx"

    # ==========================================
    # LOCAL FILE
    # ==========================================

    else:

        if not file_source:
            return ""

        extension = os.path.splitext(
            file_source
        )[1].lower()

        file_data = file_source

    # ==========================================
    # PDF
    # ==========================================

    if extension == ".pdf":

        reader = PyPDF2.PdfReader(file_data)

        for page in reader.pages:

            try:
                page_text = page.extract_text()
            except Exception:
                page_text = ""

            if page_text:
                text += page_text + "\n"

    # ==========================================
    # DOCX
    # ==========================================

    elif extension == ".docx":

        document = Document(file_data)

        for paragraph in document.paragraphs:

            if paragraph.text:
                text += paragraph.text + "\n"

        # Also extract text from tables
        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    if cell.text:
                        text += cell.text + "\n"

    # ==========================================
    # UNKNOWN FILE TYPE
    # ==========================================

    else:

        # Try PDF automatically.
        # This helps when Cloudinary URL has no extension.

        try:

            if isinstance(file_data, BytesIO):
                file_data.seek(0)

            reader = PyPDF2.PdfReader(file_data)

            for page in reader.pages:

                page_text = page.extract_text()

                if page_text:
                    text += page_text + "\n"

        except Exception:

            text = ""

    return text.strip()


# ==========================================
# RESUME ANALYSIS
# ==========================================

def analyze_resume(text):

    text_lower = text.lower()

    found_skills = []

    for skill in COMMON_SKILLS:

        if skill in text_lower:
            found_skills.append(skill)

    missing_skills = [
        skill
        for skill in COMMON_SKILLS
        if skill not in found_skills
    ]

    # ==========================================
    # ATS SCORE
    # ==========================================

    score = 50

    # Technical skills
    score += min(
        len(found_skills) * 3,
        30
    )

    # Resume length
    if len(text.split()) > 300:
        score += 10

    # ==========================================
    # EMAIL
    # ==========================================

    email_found = bool(
        re.search(
            r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
            text,
        )
    )

    # ==========================================
    # PHONE
    # ==========================================

    phone_found = bool(
        re.search(
            r"\+?\d[\d\s()-]{8,}",
            text,
        )
    )

    if email_found:
        score += 5

    if phone_found:
        score += 5

    score = min(
        score,
        100
    )

    # ==========================================
    # SUGGESTIONS
    # ==========================================

    suggestions = []

    if len(found_skills) < 6:

        suggestions.append(
            "Add more relevant technical skills."
        )

    if not email_found:

        suggestions.append(
            "Add a professional email address."
        )

    if not phone_found:

        suggestions.append(
            "Add a phone number."
        )

    if len(text.split()) < 300:

        suggestions.append(
            "Add more projects, internships, and measurable achievements."
        )

    if not suggestions:

        suggestions.append(
            "Your resume has a good overall structure. "
            "Keep improving it with measurable achievements."
        )

    return {
        "ats_score": score,
        "skills_found": found_skills,
        "missing_skills": missing_skills,
        "suggestions": suggestions,
    }